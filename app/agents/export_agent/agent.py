"""
Export Agent - Document Generation
Generates PDF, DOCX, and XLSX reports from analysis results
"""

from typing import Dict, Any, List, Optional
import logging
import json
from pathlib import Path
from io import BytesIO

logger = logging.getLogger(__name__)


class ExportAgent:
    """
    Export Agent
    
    Generates documents from analysis results:
    - PDF reports
    - DOCX documents
    - XLSX spreadsheets
    """
    
    name = "export_agent"
    supported_types = [
        "export",
        "pdf_export",
        "docx_export",
        "xlsx_export"
    ]
    
    def __init__(self):
        """Initialize Export Agent"""
        logger.info("ExportAgent initialized")
    
    def _generate_pdf(self, data: Dict[str, Any], output_path: str) -> str:
        """Generate PDF report"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import cm
            
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            # Title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(2*cm, height - 2*cm, "Stav Agent - Project Analysis Report")
            
            # Content
            c.setFont("Helvetica", 12)
            y_position = height - 4*cm
            
            # Add project summary if available
            if "project_summary" in data:
                summary = data["project_summary"]
                c.drawString(2*cm, y_position, f"Project: {summary.get('overview', 'N/A')[:50]}")
                y_position -= 0.7*cm
            
            # Add sections based on available data
            sections = [
                ("Requirements", "requirements"),
                ("Norms", "norms"),
                ("Constraints", "constraints")
            ]
            
            for section_name, section_key in sections:
                if section_key in data:
                    y_position -= 0.5*cm
                    c.drawString(2*cm, y_position, f"{section_name}:")
                    y_position -= 0.5*cm
                    
                    items = data[section_key]
                    if isinstance(items, list):
                        for item in items[:5]:  # Limit to first 5
                            c.drawString(2.5*cm, y_position, f"- {str(item)[:70]}")
                            y_position -= 0.5*cm
            
            c.save()
            return output_path
            
        except ImportError:
            logger.warning("reportlab not installed, PDF generation skipped")
            return None
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return None
    
    def _generate_docx(self, data: Dict[str, Any], output_path: str) -> str:
        """Generate DOCX document"""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('Stav Agent - Project Analysis Report', 0)
            
            # Add project summary
            if "project_summary" in data:
                doc.add_heading('Project Summary', level=1)
                summary = data["project_summary"]
                doc.add_paragraph(summary.get('overview', 'N/A'))
            
            # Add requirements
            if "requirements" in data:
                doc.add_heading('Requirements', level=1)
                for req in data["requirements"][:10]:
                    doc.add_paragraph(str(req), style='List Bullet')
            
            # Add norms
            if "norms" in data:
                doc.add_heading('Standards and Norms', level=1)
                for norm in data["norms"][:10]:
                    doc.add_paragraph(str(norm), style='List Bullet')
            
            doc.save(output_path)
            return output_path
            
        except ImportError:
            logger.warning("python-docx not installed, DOCX generation skipped")
            return None
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return None
    
    def _generate_xlsx(self, data: Dict[str, Any], output_path: str) -> str:
        """Generate XLSX spreadsheet"""
        try:
            import openpyxl
            from openpyxl import Workbook
            
            wb = Workbook()
            
            # Summary sheet
            ws = wb.active
            ws.title = "Summary"
            ws.append(["Stav Agent - Project Analysis"])
            ws.append([])
            
            if "project_summary" in data:
                ws.append(["Project Overview"])
                ws.append([data["project_summary"].get("overview", "N/A")])
                ws.append([])
            
            # Requirements sheet
            if "requirements" in data and data["requirements"]:
                ws_req = wb.create_sheet("Requirements")
                ws_req.append(["#", "Requirement"])
                for idx, req in enumerate(data["requirements"][:50], 1):
                    ws_req.append([idx, str(req)])
            
            # BOQ sheet if available
            if "items" in data and data["items"]:
                ws_boq = wb.create_sheet("BOQ")
                ws_boq.append(["#", "Code", "Description", "Quantity", "Unit", "Unit Price", "Total"])
                for item in data["items"][:100]:
                    ws_boq.append([
                        item.get("item_number", ""),
                        item.get("code", ""),
                        item.get("description", ""),
                        item.get("quantity", ""),
                        item.get("unit", ""),
                        item.get("unit_price", ""),
                        item.get("total_price", "")
                    ])
            
            wb.save(output_path)
            return output_path
            
        except ImportError:
            logger.warning("openpyxl not installed, XLSX generation skipped")
            return None
        except Exception as e:
            logger.error(f"XLSX generation failed: {e}")
            return None
    
    async def analyze(self, export_request: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate export documents.
        
        Args:
            export_request: Dictionary with:
                - data: Analysis data to export
                - format: Export format (pdf/docx/xlsx/all)
                - output_dir: Output directory path
            
        Returns:
            Dictionary with generated file paths
        """
        try:
            data = export_request.get("data", {})
            export_format = export_request.get("format", "pdf")
            output_dir = Path(export_request.get("output_dir", "/tmp"))
            
            # Ensure output directory exists
            output_dir.mkdir(parents=True, exist_ok=True)
            
            generated_files = {}
            
            # Generate based on requested format
            if export_format in ["pdf", "all"]:
                pdf_path = output_dir / "report.pdf"
                result = self._generate_pdf(data, str(pdf_path))
                if result:
                    generated_files["pdf"] = result
            
            if export_format in ["docx", "all"]:
                docx_path = output_dir / "report.docx"
                result = self._generate_docx(data, str(docx_path))
                if result:
                    generated_files["docx"] = result
            
            if export_format in ["xlsx", "all"]:
                xlsx_path = output_dir / "report.xlsx"
                result = self._generate_xlsx(data, str(xlsx_path))
                if result:
                    generated_files["xlsx"] = result
            
            result = {
                "generated_files": generated_files,
                "format": export_format,
                "status": "success" if generated_files else "partial",
                "processing_metadata": {
                    "agent": self.name,
                    "files_generated": len(generated_files)
                }
            }
            
            logger.info(f"Export completed: {len(generated_files)} files generated")
            return result
            
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return {
                "generated_files": {},
                "error": str(e),
                "processing_metadata": {
                    "agent": self.name,
                    "status": "failed"
                }
            }
    
    def supports_type(self, file_type: str) -> bool:
        """Check if this agent supports a given file type"""
        return file_type.lower() in [t.lower() for t in self.supported_types]
    
    def __repr__(self) -> str:
        """String representation of the agent"""
        return f"ExportAgent(name='{self.name}', supported_types={self.supported_types})"
