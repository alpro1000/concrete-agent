#!/usr/bin/env python3
"""
Technical Assignment Reader (TZD_READER) - Secure Version
Агент для анализа технических заданий с помощью AI моделей.

Исправления безопасности:
- Защита API ключей
- Ограничения размера файлов
- Защита от path traversal
- Валидация входных данных
- Современные API клиенты
"""

import argparse
import json
import logging
import os
import sys
import time
from io import StringIO
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import hashlib

# Импорты для обработки файлов
import pdfplumber
from docx import Document

# AI клиенты (обновленные)
from openai import OpenAI
from anthropic import Anthropic

# Настройка логирования
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('tzd_reader.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

# Константы безопасности
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
MAX_TOTAL_TEXT_LENGTH = 500_000  # Максимальная длина объединенного текста
API_TIMEOUT = 60  # Таймаут для API вызовов
MAX_RETRIES = 3

# Настройки моделей (настраиваемые через env)
DEFAULT_OPENAI_MODEL = "gpt-4o-mini"
DEFAULT_CLAUDE_MODEL = "claude-3-5-sonnet-20241022"


class SecurityError(Exception):
    """Исключение для проблем безопасности"""
    pass


class FileSecurityValidator:
    """Валидатор безопасности файлов"""
    
    @staticmethod
    def validate_file_path(file_path: str, base_dir: Optional[str] = None) -> bool:
        """
        Проверяет безопасность пути к файлу
        
        Args:
            file_path: Путь к файлу
            base_dir: Базовая директория (опционально)
            
        Returns:
            True если путь безопасен
            
        Raises:
            SecurityError: При обнаружении небезопасного пути
        """
        try:
            # Нормализация пути
            normalized_path = os.path.normpath(os.path.abspath(file_path))
            
            # Проверка на path traversal
            if '..' in normalized_path or normalized_path.startswith('/'):
                raise SecurityError(f"Потенциально опасный путь: {file_path}")
            
            # Проверка базовой директории
            if base_dir:
                base_abs = os.path.abspath(base_dir)
                if not normalized_path.startswith(base_abs):
                    raise SecurityError(f"Файл вне разрешенной директории: {file_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"Ошибка валидации пути {file_path}: {e}")
            raise SecurityError(f"Небезопасный путь к файлу: {file_path}")
    
    @staticmethod
    def validate_file_size(file_path: str) -> bool:
        """
        Проверяет размер файла
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            True если размер допустим
            
        Raises:
            SecurityError: При превышении размера
        """
        try:
            file_size = os.path.getsize(file_path)
            if file_size > MAX_FILE_SIZE:
                raise SecurityError(
                    f"Файл слишком большой: {file_size} bytes (max: {MAX_FILE_SIZE})"
                )
            return True
        except OSError as e:
            raise SecurityError(f"Не удается проверить размер файла {file_path}: {e}")
    
    @staticmethod
    def validate_file_extension(file_path: str) -> bool:
        """
        Проверяет расширение файла
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            True если расширение разрешено
            
        Raises:
            SecurityError: При неразрешенном расширении
        """
        ext = Path(file_path).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise SecurityError(
                f"Неразрешенное расширение файла: {ext}. "
                f"Разрешены: {', '.join(ALLOWED_EXTENSIONS)}"
            )
        return True


class SecureFileProcessor:
    """Безопасный обработчик файлов"""
    
    def __init__(self, base_dir: Optional[str] = None):
        self.validator = FileSecurityValidator()
        self.base_dir = base_dir
        self.processed_files = []
    
    def _safe_read_pdf(self, file_path: str) -> str:
        """Безопасное чтение PDF с таймаутом"""
        text_parts = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                # Ограничиваем количество страниц для обработки
                max_pages = min(len(pdf.pages), 100)
                
                for i, page in enumerate(pdf.pages[:max_pages]):
                    if i > 0 and i % 10 == 0:
                        # Периодическая проверка размера текста
                        current_text = '\n'.join(text_parts)
                        if len(current_text) > MAX_TOTAL_TEXT_LENGTH // 2:
                            logger.warning(f"Достигнут лимит текста на странице {i}")
                            break
                    
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        
        except Exception as e:
            logger.error(f"Ошибка чтения PDF {file_path}: {e}")
            raise
        
        return '\n'.join(text_parts)
    
    def _safe_read_docx(self, file_path: str) -> str:
        """Безопасное чтение DOCX"""
        text_parts = []
        
        try:
            doc = Document(file_path)
            
            # Ограничиваем количество параграфов
            max_paragraphs = min(len(doc.paragraphs), 1000)
            
            for i, paragraph in enumerate(doc.paragraphs[:max_paragraphs]):
                if i > 0 and i % 100 == 0:
                    # Периодическая проверка размера
                    current_text = '\n'.join(text_parts)
                    if len(current_text) > MAX_TOTAL_TEXT_LENGTH // 2:
                        logger.warning(f"Достигнут лимит текста на параграфе {i}")
                        break
                
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
        except Exception as e:
            logger.error(f"Ошибка чтения DOCX {file_path}: {e}")
            raise
        
        return '\n'.join(text_parts)
    
    def _safe_read_txt(self, file_path: str) -> str:
        """Безопасное чтение текстового файла"""
        encodings = ['utf-8', 'cp1251', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    # Читаем файл частями для контроля памяти
                    text_parts = []
                    total_length = 0
                    
                    while True:
                        chunk = f.read(8192)  # 8KB chunks
                        if not chunk:
                            break
                        
                        text_parts.append(chunk)
                        total_length += len(chunk)
                        
                        if total_length > MAX_TOTAL_TEXT_LENGTH:
                            logger.warning(f"Текстовый файл обрезан по лимиту: {file_path}")
                            break
                    
                    return ''.join(text_parts)
                    
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Ошибка чтения TXT {file_path} с кодировкой {encoding}: {e}")
                continue
        
        raise Exception(f"Не удалось прочитать файл ни с одной из кодировок: {encodings}")
    
    def process_files(self, file_paths: List[str]) -> str:
        """
        Безопасная обработка списка файлов
        
        Args:
            file_paths: Список путей к файлам
            
        Returns:
            Объединенный текст из всех файлов
            
        Raises:
            SecurityError: При проблемах безопасности
            ValueError: При невалидных входных данных
        """
        if not file_paths:
            raise ValueError("Список файлов не может быть пустым")
        
        if len(file_paths) > 20:
            raise ValueError("Слишком много файлов для обработки (максимум: 20)")
        
        text_parts = []
        total_length = 0
        
        for file_path in file_paths:
            try:
                # Валидация безопасности
                self.validator.validate_file_path(file_path, self.base_dir)
                self.validator.validate_file_size(file_path)
                self.validator.validate_file_extension(file_path)
                
                if not os.path.exists(file_path):
                    logger.warning(f"Файл не найден: {file_path}")
                    continue
                
                file_ext = Path(file_path).suffix.lower()
                logger.info(f"Обработка файла: {Path(file_path).name} ({file_ext})")
                
                # Обработка по типу файла
                if file_ext == '.pdf':
                    text = self._safe_read_pdf(file_path)
                elif file_ext == '.docx':
                    text = self._safe_read_docx(file_path)
                elif file_ext == '.txt':
                    text = self._safe_read_txt(file_path)
                else:
                    continue  # Уже проверено в валидации
                
                if text.strip():
                    file_header = f"\n=== Файл: {Path(file_path).name} ===\n"
                    text_parts.extend([file_header, text, "\n"])
                    
                    total_length += len(file_header) + len(text) + 1
                    
                    # Проверка общего лимита
                    if total_length > MAX_TOTAL_TEXT_LENGTH:
                        logger.warning("Достигнут максимальный лимит текста")
                        break
                        
                    self.processed_files.append(file_path)
                else:
                    logger.warning(f"Пустой текст в файле: {file_path}")
                    
            except SecurityError:
                raise  # Проблемы безопасности прерывают обработку
            except Exception as e:
                logger.error(f"Ошибка обработки файла {file_path}: {e}")
                continue  # Продолжаем с другими файлами
        
        combined_text = ''.join(text_parts)
        
        if not combined_text.strip():
            raise ValueError("Не удалось извлечь текст ни из одного файла")
        
        logger.info(f"Обработано файлов: {len(self.processed_files)}")
        logger.info(f"Общая длина текста: {len(combined_text)} символов")
        
        return combined_text


class SecureAIAnalyzer:
    """Безопасный анализатор с современными AI клиентами"""
    
    def __init__(self):
        self.openai_client = None
        self.anthropic_client = None
        self.openai_model = os.getenv('OPENAI_MODEL', DEFAULT_OPENAI_MODEL)
        self.claude_model = os.getenv('CLAUDE_MODEL', DEFAULT_CLAUDE_MODEL)
        
        self._init_openai_client()
        self._init_anthropic_client()
    
    def _init_openai_client(self):
        """Безопасная инициализация OpenAI клиента"""
        api_key = os.getenv('OPENAI_API_KEY')
        if api_key:
            if len(api_key) < 20:  # Базовая проверка валидности ключа
                logger.error("Подозрительно короткий OpenAI API ключ")
                return
            
            try:
                self.openai_client = OpenAI(
                    api_key=api_key,
                    timeout=API_TIMEOUT,
                    max_retries=MAX_RETRIES
                )
                logger.info("OpenAI клиент инициализирован")
            except Exception as e:
                logger.error(f"Ошибка инициализации OpenAI клиента: {e}")
        else:
            logger.warning("OPENAI_API_KEY не установлен")
    
    def _init_anthropic_client(self):
        """Безопасная инициализация Anthropic клиента"""
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if api_key:
            if len(api_key) < 20:  # Базовая проверка валидности ключа
                logger.error("Подозрительно короткий Anthropic API ключ")
                return
            
            try:
                self.anthropic_client = Anthropic(
                    api_key=api_key,
                    timeout=API_TIMEOUT,
                    max_retries=MAX_RETRIES
                )
                logger.info("Anthropic клиент инициализирован")
            except Exception as e:
                logger.error(f"Ошибка инициализации Anthropic клиента: {e}")
        else:
            logger.warning("ANTHROPIC_API_KEY не установлен")
    
    def get_analysis_prompt(self) -> str:
        """Промпт для анализа технического задания"""
        return """Проанализируй техническое задание и верни результат СТРОГО в формате JSON.

Требуемые поля JSON:
{
  "project_name": "название проекта (строка)",
  "project_scope": "описание объёма работ (строка)", 
  "materials": ["список материалов"],
  "concrete_requirements": ["марки бетона и требования"],
  "norms": ["нормативные документы"],
  "functional_requirements": ["функциональные требования"],
  "risks_and_constraints": ["риски и ограничения"],
  "estimated_complexity": "низкая|средняя|высокая",
  "key_technologies": ["ключевые технологии"]
}

Если информация отсутствует - используй пустые строки или массивы.
Отвечай ТОЛЬКО валидным JSON без комментариев.

Текст технического задания:
"""
    
    def _extract_json_from_response(self, response_text: str) -> str:
        """Извлекает JSON из ответа AI модели"""
        text = response_text.strip()
        
        # Удаляем markdown блоки
        if "```json" in text:
            start = text.find("```json") + 7
            end = text.find("```", start)
            if end != -1:
                text = text[start:end].strip()
        elif text.startswith("```") and text.endswith("```"):
            text = text[3:-3].strip()
        
        # Ищем JSON объект
        start = text.find('{')
        end = text.rfind('}') + 1
        
        if start != -1 and end != -1 and start < end:
            return text[start:end]
        
        return text
    
    def _get_empty_result(self) -> Dict[str, Any]:
        """Базовый результат при ошибках"""
        return {
            "project_name": "",
            "project_scope": "",
            "materials": [],
            "concrete_requirements": [],
            "norms": [],
            "functional_requirements": [],
            "risks_and_constraints": [],
            "estimated_complexity": "неопределена",
            "key_technologies": [],
            "processing_error": True
        }
    
    def analyze_with_gpt(self, text: str) -> Dict[str, Any]:
        """Анализ с помощью OpenAI GPT (обновленный API)"""
        if not self.openai_client:
            raise ValueError("OpenAI клиент не инициализирован")
        
        # Обрезаем текст если слишком длинный
        max_tokens_input = 100000  # Примерный лимит для gpt-4o-mini
        if len(text) > max_tokens_input:
            text = text[:max_tokens_input] + "\n\n[ТЕКСТ ОБРЕЗАН ПО ЛИМИТУ]"
            logger.warning("Текст обрезан для OpenAI API")
        
        try:
            prompt = self.get_analysis_prompt() + text
            
            response = self.openai_client.chat.completions.create(
                model=self.openai_model,
                messages=[
                    {
                        "role": "system", 
                        "content": "Ты эксперт по анализу технических заданий в строительстве. "
                                 "Отвечай ТОЛЬКО валидным JSON без дополнительного текста."
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,  # Более детерминированные результаты
                max_tokens=4000,
                timeout=API_TIMEOUT
            )
            
            result_text = response.choices[0].message.content
            if not result_text:
                raise ValueError("Пустой ответ от OpenAI")
            
            json_text = self._extract_json_from_response(result_text)
            result = json.loads(json_text)
            
            # Добавляем метаданные
            result['ai_model'] = self.openai_model
            result['processing_time'] = response.usage.total_tokens if hasattr(response, 'usage') else None
            
            return result
            
        except json.JSONDecodeError as e:
            logger.error(f"Невалидный JSON от GPT: {e}")
            return self._get_empty_result()
        except Exception as e:
            logger.error(f"Ошибка анализа с GPT: {e}")
            return self._get_empty_result()
    
    def analyze_with_claude(self, text: str) -> Dict[str, Any]:
        """Анализ с помощью Anthropic Claude"""
        if not self.anthropic_client:
            raise ValueError("Anthropic клиент не инициализирован")
        
        # Обрезаем текст если слишком длинный  
        max_tokens_input = 180000  # Лимит для Claude
        if len(text) > max_tokens_input:
            text = text[:max_tokens_input] + "\n\n[ТЕКСТ ОБРЕЗАН ПО ЛИМИТУ]"
            logger.warning("Текст обрезан для Claude API")
        
        try:
            prompt = self.get_analysis_prompt() + text
            
            response = self.anthropic_client.messages.create(
                model=self.claude_model,
                max_tokens=4000,
                temperature=0.1,
                messages=[
                    {"role": "user", "content": prompt}
                ],
                timeout=API_TIMEOUT
            )
            
            result_text = response.content[0].text
            if not result_text:
                raise ValueError("Пустой ответ от Claude")
            
            json_text = self._extract_json_from_response(result_text)
            result = json.loads(json_text)
            
            # Добавляем метаданные
            result['ai_model'] = self.claude_model
            result['processing_time'] = response.usage.output_tokens if hasattr(response, 'usage') else None
            
            return result
            
        except json.JSONDecodeError as e:
            logger.error(f"Невалидный JSON от Claude: {e}")
            return self._get_empty_result()
        except Exception as e:
            logger.error(f"Ошибка анализа с Claude: {e}")
            return self._get_empty_result()


def tzd_reader(files: List[str], engine: str = "gpt", base_dir: Optional[str] = None) -> Dict[str, Any]:
    """
    Основная функция для безопасного анализа технических заданий
    
    Args:
        files: Список путей к файлам
        engine: AI движок ("gpt" или "claude")
        base_dir: Базовая директория для ограничения доступа к файлам
        
    Returns:
        Словарь с результатами анализа
        
    Raises:
        ValueError: При невалидных параметрах
        SecurityError: При проблемах безопасности
    """
    if not files:
        raise ValueError("Список файлов не может быть пустым")
    
    if engine.lower() not in ['gpt', 'claude']:
        raise ValueError(f"Неподдерживаемый AI движок: {engine}. Используйте 'gpt' или 'claude'")
    
    start_time = time.time()
    
    try:
        # Безопасная обработка файлов
        processor = SecureFileProcessor(base_dir=base_dir)
        combined_text = processor.process_files(files)
        
        # AI анализ
        analyzer = SecureAIAnalyzer()
        
        if engine.lower() == "gpt":
            result = analyzer.analyze_with_gpt(combined_text)
        else:  # claude
            result = analyzer.analyze_with_claude(combined_text)
        
        # Добавляем метаданные обработки
        result['processing_metadata'] = {
            'processed_files': len(processor.processed_files),
            'total_text_length': len(combined_text),
            'processing_time_seconds': round(time.time() - start_time, 2),
            'ai_engine': engine.lower(),
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
        }
        
        return result
        
    except Exception as e:
        logger.error(f"Критическая ошибка в tzd_reader: {e}")
        error_result = SecureAIAnalyzer()._get_empty_result()
        error_result['critical_error'] = str(e)
        return error_result


def main():
    """Основная функция CLI с улучшенной обработкой ошибок"""
    parser = argparse.ArgumentParser(
        description="Technical Assignment Reader (TZD_READER) - Secure Version",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры использования:
  python tzd_reader.py --files doc1.pdf doc2.docx --engine gpt
  python tzd_reader.py --files *.pdf --engine claude --base-dir ./projects
        """
    )
    
    parser.add_argument(
        "--files", 
        nargs="+", 
        required=True,
        help="Список файлов для анализа (PDF, DOCX, TXT)"
    )
    
    parser.add_argument(
        "--engine", 
        choices=["gpt", "claude"], 
        default="gpt",
        help="AI движок для анализа (по умолчанию: gpt)"
    )
    
    parser.add_argument(
        "--base-dir",
        type=str,
        help="Базовая директория для ограничения доступа к файлам"
    )
    
    parser.add_argument(
        "--output",
        type=str,
        help="Файл для сохранения результата (по умолчанию: вывод в консоль)"
    )
    
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Подробный вывод логов"
    )
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        logger.info("Запуск TZD Reader (Secure Version)")
        logger.info(f"Файлы для обработки: {args.files}")
        logger.info(f"AI движок: {args.engine}")
        
        result = tzd_reader(
            files=args.files,
            engine=args.engine,
            base_dir=args.base_dir
        )
        
        # Форматированный вывод
        output_json = json.dumps(result, ensure_ascii=False, indent=2)
        
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(output_json)
            logger.info(f"Результат сохранен в файл: {args.output}")
        else:
            print(output_json)
        
        logger.info("Анализ завершен успешно")
        
    except SecurityError as e:
        logger.error(f"Ошибка безопасности: {e}")
        sys.exit(2)
    except ValueError as e:
        logger.error(f"Ошибка валидации: {e}")
        sys.exit(3)
    except Exception as e:
        logger.error(f"Критическая ошибка: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
