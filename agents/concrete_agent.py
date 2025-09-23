# parsers/doc_parser.py - ИСПРАВЛЕННАЯ ВЕРСИЯ без textract
"""
Парсер документов без использования textract
Замена для оригинального doc_parser.py
"""
import os
import logging
from pathlib import Path
from typing import Optional, Dict, List

logger = logging.getLogger(__name__)

def extract_text_from_document(file_path: str) -> Optional[str]:
    """
    Извлечение текста из документов без textract
    Поддерживает: PDF, DOCX, TXT
    """
    try:
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.txt':
            return _extract_from_txt(file_path)
        elif extension == '.pdf':
            return _extract_from_pdf(file_path)
        elif extension in ['.doc', '.docx']:
            return _extract_from_docx(file_path)
        else:
            logger.warning(f"Неподдерживаемый формат: {extension}")
            return None
            
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из {file_path}: {e}")
        return None

def _extract_from_txt(file_path: Path) -> Optional[str]:
    """Извлечение из TXT файлов"""
    encodings = ['utf-8', 'cp1251', 'iso-8859-1', 'windows-1252']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    
    logger.error(f"Не удалось определить кодировку для {file_path}")
    return None

def _extract_from_pdf(file_path: Path) -> Optional[str]:
    """Извлечение из PDF файлов"""
    try:
        # Используем pdfplumber (лучший выбор)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        except ImportError:
            logger.warning("pdfplumber не установлен")
        
        # Fallback к PyPDF2
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except ImportError:
            logger.warning("PyPDF2 не установлен")
            
        logger.error("Нет библиотек для работы с PDF")
        return None
        
    except Exception as e:
        logger.error(f"Ошибка чтения PDF: {e}")
        return None

def _extract_from_docx(file_path: Path) -> Optional[str]:
    """Извлечение из DOCX файлов"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
        
    except ImportError:
        logger.error("python-docx не установлен")
        return None
    except Exception as e:
        logger.error(f"Ошибка чтения DOCX: {e}")
        return None

# Совместимость с оригинальным API
def parse_document(file_path: str) -> Dict:
    """
    Основная функция парсинга документа
    Возвращает структурированный результат
    """
    text = extract_text_from_document(file_path)
    
    if not text:
        return {
            "success": False,
            "error": f"Не удалось извлечь текст из {file_path}",
            "text": "",
            "metadata": {}
        }
    
    # Базовая обработка текста
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    word_count = len(text.split())
    
    return {
        "success": True,
        "text": text,
        "metadata": {
            "file_path": file_path,
            "lines_count": len(lines),
            "word_count": word_count,
            "char_count": len(text)
        }
    }

def get_supported_formats() -> List[str]:
    """Возвращает список поддерживаемых форматов"""
    return ['.txt', '.pdf', '.docx', '.doc']

# Проверка доступности библиотек
def check_dependencies() -> Dict[str, bool]:
    """Проверяет доступность необходимых библиотек"""
    deps = {
        'pdfplumber': False,
        'PyPDF2': False,
        'python-docx': False
    }
    
    for lib in deps:
        try:
            if lib == 'python-docx':
                import docx
            else:
                __import__(lib)
            deps[lib] = True
        except ImportError:
            pass
    
    return deps

if __name__ == "__main__":
    # Тестирование
    print("🔍 Проверка зависимостей:")
    deps = check_dependencies()
    for lib, available in deps.items():
        status = "✅" if available else "❌"
        print(f"{status} {lib}")
    
    print(f"\n📄 Поддерживаемые форматы: {', '.join(get_supported_formats())}")
